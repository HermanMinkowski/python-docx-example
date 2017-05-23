#!/usr/bin/env python

"""
create_docx.py: Short demonstration how to use python-docx (http://python-docx.readthedocs.io)
to generate word documents automatically.

Dependencies:
 -  Python 3.5 (should work with other versions)
 -  lxml >= 2.3.2
 -  python-docx >= 0.8.6
"""

from json import load
from datetime import date
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

__author__ = "Pascal Renauld"
__copyright__ = "Copyright 2017"
__license__ = "WTFPL"
__version__ = "0.1.1.0.1.1.1.1.0.1.1.0.1.1.1.0.0.1.1.0.0.1.0.1"


with open('ransomware_overview.json') as json_file:
    REPORT = load(json_file)

document = Document("Template.docx")

# Setting the document core properties (used for the footer).
document.core_properties.author = "Jean Thomas Dulaien"

# We are using the default word styles, level 0 is Title.
# Level 1 to 9 are Heading X.
document.add_heading("List of ransomware", level=0)

# This adds an inline picture in its own paragraph.
# Word measures things in EMU there are 914,400 to the inch.
# You can use docx.shared to convert other units to EMUs.
picture = document.add_picture("Captain_Dulaien_Flag.png", Inches(2.5))

# This is a simple paragraph. You can apply a style to this paragraph.
paragraph = document.add_paragraph("This ransomware report has been generated for ", "Pirate")

# Runs are elements in a paragraph. You can format runs using italic, bold, underline etc.
paragraph.add_run("Captain Dulaien ").bold = True
paragraph.add_run("on "  + str(date.today()) + ".")

# Adds a heading 1.
document.add_heading("List of ransomware", level=1)

# Adds a heading and a table for each ransomware.
for ranswomware in REPORT:
    document.add_heading(ranswomware["name"], level=2)

    # Creates a table.
    table = document.add_table(rows=4, cols=2)

    # Applies the Report style to the whole table.
    table.style = "Report"

    # Adds text to the first cell of the first row
    table.rows[0].cells[0].text = "extensions"
    # Changes the property of the paragraph for this cell to add the "keep with next feature".
    # This keep all the table lines together.
    table.rows[0].cells[0].paragraphs[0].paragraph_format.keep_with_next = True

    table.rows[0].cells[1].text = ranswomware["extensions"]
    table.rows[0].cells[1].paragraphs[0].paragraph_format.keep_with_next = True

    table.rows[1].cells[0].text = "ransomNoteFilenames"
    table.rows[1].cells[0].paragraphs[0].paragraph_format.keep_with_next = True
    table.rows[1].cells[1].text = ranswomware["ransomNoteFilenames"]
    table.rows[1].cells[1].paragraphs[0].paragraph_format.keep_with_next = True

    table.rows[2].cells[0].text = "encryptionAlgorithm"
    table.rows[2].cells[0].paragraphs[0].paragraph_format.keep_with_next = True
    table.rows[2].cells[1].text = ranswomware["encryptionAlgorithm"]
    table.rows[2].cells[1].paragraphs[0].paragraph_format.keep_with_next = True

    table.rows[3].cells[0].text = "comment"
    table.rows[3].cells[1].text = ranswomware["comment"]


# To center a picture or gain more control on it you can first
# create an empty paragraph, then a run, then add a picture to the run.
paragraph_with_picture = document.add_paragraph("END OF FILE - ")
picture_run = paragraph_with_picture.add_run()
picture_run.add_picture("Captain_Dulaien_Flag.png", Inches(0.5))
paragraph_with_picture.add_run(" - END OF FILE")
paragraph_with_picture.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

# If your docx is large, it can take several seconds/minutes to generate
document.save("ransomware_report.docx")




